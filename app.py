import streamlit as st
import os
import json
import requests
import numpy as np
import faiss
import time
import pickle
from github import Github, GithubException, Auth  # 新增 Auth
from pathlib import Path
from io import BytesIO
from typing import List, Dict, Any, Tuple, Optional
from http import HTTPStatus
import dashscope
from dashscope import TextEmbedding
from openai import OpenAI
from docx import Document
import matplotlib.pyplot as plt
from scipy.interpolate import make_interp_spline
from graphrag_retriever import GraphRAGRetriever, GraphRAGIndexer, Chunk
import base64
import hashlib
import fitz          # PyMuPDF: 渲染PDF为图片 + 文本提取
import pytesseract   # OCR引擎接口
from PIL import Image

# ==========================================
# [SECTION 0] 基础配置与路径定义
# ==========================================

st.set_page_config(
    page_title="茶饮六因子AI评分器 Pro",
    page_icon="🍵",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 样式定义
st.markdown("""
    <style>
    .main-title {font-size: 2.5em; font-weight: bold; text-align: center; color: #2E7D32; margin-bottom: 0.5em;}
    .slogan {font-size: 1.2em; font-style: italic; text-align: center; color: #558B2F; margin-bottom: 30px; font-family: "KaiTi", "楷体", serif;}
    .factor-card {background-color: #F1F8E9; padding: 15px; border-radius: 10px; margin-bottom: 10px; border-left: 5px solid #4CAF50;}
    .score-header {display:flex; justify-content:space-between; font-weight:bold; color:#2E7D32;}
    .advice-tag {font-size: 0.85em; padding: 2px 6px; border-radius: 4px; margin-top: 5px; background-color: #fff; border: 1px dashed #4CAF50; color: #388E3C; display: inline-block;}
    .master-comment {background-color: #FFFDE7; border: 1px solid #FFF9C4; padding: 15px; border-radius: 8px; font-family: "KaiTi", serif; font-size: 1.1em; color: #5D4037; margin-bottom: 20px; line-height: 1.6;}
    .ft-card {border: 1px solid #ddd; padding: 15px; border-radius: 8px; background-color: #f8f9fa; margin-top: 10px;}
    .case-card {border: 1px solid #e0e0e0; padding: 12px; border-radius: 8px; margin-bottom: 10px; background-color: #fafafa;}
    </style>
""", unsafe_allow_html=True)

class PathConfig:
    """路径管理类"""
    # 外部资源文件（位于同级目录）
    SRC_SYS_PROMPT = Path("sys_p.txt")
    SRC_CASES = Path("tea_data/case.json")  # Case文件存储
    # 运行时数据目录
    DATA_DIR = Path("./tea_data")
    RAG_DIR = Path("./tea_data/RAG")  # RAG文件存储目录
    
    def __init__(self):
        self.DATA_DIR.mkdir(exist_ok=True)
        self.RAG_DIR.mkdir(exist_ok=True)  # 确保RAG目录存在
        self.GRAPHRAG_DIR = self.DATA_DIR / "graphrag_artifacts"
        self.GRAPHRAG_DIR.mkdir(exist_ok=True)
        self.OCR_CACHE_DIR = self.DATA_DIR / "ocr_cache"
        self.OCR_CACHE_DIR.mkdir(exist_ok=True)
        # 向量库与持久化数据
        self.kb_index = self.DATA_DIR / "kb.index"
        self.kb_chunks = self.DATA_DIR / "kb_chunks.pkl"
        self.kb_files = self.DATA_DIR / "kb_files.json"  # 新增：记录RAG文件列表
        self.case_index = self.DATA_DIR / "cases.index"
        self.case_data = self.DATA_DIR / "case.json"  # 修改：与GitHub保持一致
        
        # 微调与Prompt配置
        self.training_file = self.DATA_DIR / "deepseek_finetune.jsonl"
        self.ft_status = self.DATA_DIR / "ft_status.json"
        self.prompt_config_file = self.DATA_DIR / "prompts.json"

PATHS = PathConfig()

# 默认的用户Prompt模板（System Prompt将从文件读取）
DEFAULT_USER_TEMPLATE = """【待评分产品】
{product_desc}

【参考标准（知识库）】
{context_text}

【相似判例得分参考（案例库）】
{case_text}

请严格输出以下JSON格式（不含Markdown）：
{{
  "master_comment": "约100字的宗师级总评，富含文化意蕴...",
  "scores": {{
    "优雅性": {{"score": 0-9, "comment": "...", "suggestion": "..."}},
    "辨识度": {{"score": 0-9, "comment": "...", "suggestion": "..."}},
    "协调性": {{"score": 0-9, "comment": "...", "suggestion": "..."}},
    "饱和度": {{"score": 0-9, "comment": "...", "suggestion": "..."}},
    "持久性": {{"score": 0-9, "comment": "...", "suggestion": "..."}},
    "苦涩度": {{"score": 0-9, "comment": "...", "suggestion": "..."}}
  }}
}}"""

# ==========================================
# [SECTION 1] 资源与数据管理
# ==========================================

class ResourceManager:
    """负责外部文件加载、数据持久化及格式转换"""

    @staticmethod
    def load_external_text(path: Path, fallback: str = "") -> str:
        """读取外部文本文件 (如 sys_p.txt)"""
        if path.exists():
            try:
                with open(path, "r", encoding="utf-8") as f:
                    return f.read().strip()
            except Exception as e:
                st.error(f"加载文件 {path} 失败: {e}")
        return fallback

    @staticmethod
    def load_external_json(path: Path, fallback: Any = None) -> Any:
        """读取外部JSON文件"""
        if path.exists():
            try:
                with open(path, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception as e:
                st.error(f"加载文件 {path} 失败: {e}")
        return fallback if fallback is not None else []

    @staticmethod
    def save(index: Any, data: Any, idx_path: Path, data_path: Path, is_json: bool = False):
        """保存 FAISS 索引和数据文件"""
        if index: faiss.write_index(index, str(idx_path))
        with open(data_path, "w" if is_json else "wb", encoding="utf-8" if is_json else None) as f:
            if is_json: json.dump(data, f, ensure_ascii=False, indent=2)
            else: pickle.dump(data, f)
    
    @staticmethod
    def load(idx_path: Path, data_path: Path, is_json: bool = False) -> Tuple[Any, List]:
        """加载 FAISS 索引和数据文件"""
        if idx_path.exists() and data_path.exists():
            try:
                index = faiss.read_index(str(idx_path))
                with open(data_path, "r" if is_json else "rb", encoding="utf-8" if is_json else None) as f:
                    data = json.load(f) if is_json else pickle.load(f)
                return index, data
            except: pass
        return faiss.IndexFlatL2(1024), []

    # ===== 微调相关方法 =====
    @staticmethod
    def overwrite_finetune(cases: List[Dict], sys_prompt: str, user_tpl: str) -> int:
        """覆盖写入微调数据集 (.jsonl) - 修改为覆盖逻辑"""
        try:
            count = 0
            with open(PATHS.training_file, "w", encoding="utf-8") as f:
                for c in cases:
                    case_text = c.get("text", "")
                    scores = c.get("scores", {})
                    master_comment = c.get("master_comment", "（人工校准）")
                    
                    user_content = user_tpl.replace("{product_desc}", case_text).replace("{context_text}", "").replace("{case_text}", "")
                    assistant_content = json.dumps({"master_comment": master_comment, "scores": scores}, ensure_ascii=False)
                    entry = {
                        "messages": [
                            {"role": "system", "content": sys_prompt},
                            {"role": "user", "content": user_content},
                            {"role": "assistant", "content": assistant_content}
                        ]
                    }
                    f.write(json.dumps(entry, ensure_ascii=False) + "\n")
                    count += 1
            return count
        except Exception as e:
            print(f"[ERROR] Finetune overwrite failed: {e}")
            return 0

    @staticmethod
    def save_ft_status(job_id, status, fine_tuned_model=None):
        data = {"job_id": job_id, "status": status, "timestamp": time.time()}
        if fine_tuned_model: data["fine_tuned_model"] = fine_tuned_model
        with open(PATHS.ft_status, 'w') as f: json.dump(data, f)

    @staticmethod
    def load_ft_status():
        if PATHS.ft_status.exists():
            try: return json.load(open(PATHS.ft_status, 'r'))
            except: pass
        return None

    # ===== RAG文件管理 =====
    @staticmethod
    def save_kb_files(file_list: List[str]):
        """保存知识库文件列表"""
        with open(PATHS.kb_files, "w", encoding="utf-8") as f:
            json.dump(file_list, f, ensure_ascii=False, indent=2)
    
    @staticmethod
    def load_kb_files() -> List[str]:
        """加载知识库文件列表"""
        if PATHS.kb_files.exists():
            try:
                with open(PATHS.kb_files, "r", encoding="utf-8") as f:
                    return json.load(f)
            except: pass
        return []

# ==========================================
# [SECTION 1.2] OCR 缓存管理
# ==========================================

class OCRCache:
    """
    PDF OCR 缓存管理。
    对 PDF content 计算 MD5 作为唯一 key，OCR 结果存为 .txt。
    本地缓存: tea_data/ocr_cache/{md5}.txt
    GitHub 缓存: tea_data/ocr_cache/{md5}.txt（持久化，跨部署生效）
    """

    @staticmethod
    def _md5(content: bytes) -> str:
        return hashlib.md5(content).hexdigest()

    @staticmethod
    def get(content: bytes) -> str | None:
        """查缓存，命中返回文本，未命中返回 None"""
        md5 = OCRCache._md5(content)
        cache_path = PATHS.OCR_CACHE_DIR / f"{md5}.txt"
        if cache_path.exists():
            try:
                text = cache_path.read_text(encoding="utf-8")
                if text.strip():
                    print(f"[INFO]     → OCR 缓存命中: {md5[:8]}... ({len(text):,} 字符)")
                    return text
            except:
                pass
        return None

    @staticmethod
    def put(content: bytes, text: str, push_to_github: bool = True):
        """写缓存（本地 + GitHub）"""
        md5 = OCRCache._md5(content)
        cache_path = PATHS.OCR_CACHE_DIR / f"{md5}.txt"
        cache_path.write_text(text, encoding="utf-8")
        print(f"[INFO]     → OCR 结果已缓存: {md5[:8]}... ({len(text):,} 字符)")
        if push_to_github:
            try:
                github_path = f"tea_data/ocr_cache/{md5}.txt"
                GithubSync.push_binary_file(github_path, text.encode("utf-8"), f"Cache OCR: {md5[:8]}")
                print(f"[INFO]     → 缓存已同步到 GitHub")
            except Exception as e:
                print(f"[WARN]     → GitHub 缓存同步失败（不影响使用）: {e}")

    @staticmethod
    def pull_all_from_github():
        """启动时从 GitHub 拉取所有 OCR 缓存到本地（仅本地缓存为空时执行）"""
        try:
            local_files = list(PATHS.OCR_CACHE_DIR.glob("*.txt"))
            if local_files:
                print(f"[INFO] 本地已有 {len(local_files)} 个 OCR 缓存，跳过 GitHub 拉取")
                return
            
            files = GithubSync.pull_rag_folder("tea_data/ocr_cache")
            count = 0
            for fname, content in files:
                if fname.endswith(".txt"):
                    local_path = PATHS.OCR_CACHE_DIR / fname
                    if not local_path.exists():
                        local_path.write_bytes(content)
                        count += 1
            if count > 0:
                print(f"[INFO] 从 GitHub 拉取了 {count} 个 OCR 缓存文件")
        except Exception as e:
            print(f"[WARN] 拉取 OCR 缓存失败（不影响使用）: {e}")

# ==========================================
# [SECTION 1.5] Github 同步工具 (增强版)
# ==========================================

class GithubSync:
    """负责将数据同步回 Github 仓库"""
    
    @staticmethod
    def _get_github_config():
        """获取GitHub配置"""
        token = st.secrets.get("GITHUB_TOKEN")
        repo_name = st.secrets.get("GITHUB_REPO")
        branch = st.secrets.get("GITHUB_BRANCH", "main")
        return token, repo_name, branch
    
    @staticmethod
    def _get_github_client():
        """获取 GitHub 客户端（使用新的认证方式）"""
        token, repo_name, branch = GithubSync._get_github_config()
        if not token or not repo_name:
            return None, None, None
        # 使用新的认证方式，避免 DeprecationWarning
        g = Github(auth=Auth.Token(token))
        return g, repo_name, branch
    
    @staticmethod
    def push_json(file_path_in_repo: str, data_dict: Dict, commit_msg: str = "Update via Streamlit") -> bool:
        """推送 JSON 数据到 Github"""
        g, repo_name, branch = GithubSync._get_github_client()
        
        if not g or not repo_name:
            st.error("❌ 未配置 Github Token 或 仓库名 (GITHUB_TOKEN / GITHUB_REPO)")
            return False

        try:
            repo = g.get_repo(repo_name)
            content_str = json.dumps(data_dict, ensure_ascii=False, indent=2)
            
            try:
                contents = repo.get_contents(file_path_in_repo, ref=branch)
                repo.update_file(
                    path=contents.path,
                    message=commit_msg,
                    content=content_str,
                    sha=contents.sha,
                    branch=branch
                )
            except GithubException as e:
                if e.status == 404:
                    repo.create_file(
                        path=file_path_in_repo,
                        message=f"Create {file_path_in_repo}",
                        content=content_str,
                        branch=branch
                    )
                else:
                    raise e
            return True

        except Exception as e:
            st.error(f"Github 同步失败: {str(e)}")
            return False

    @staticmethod
    def push_binary_file(file_path_in_repo: str, file_content: bytes, commit_msg: str = "Upload file") -> bool:
        """推送二进制文件到 Github (如PDF, DOCX等)
        
        重要：PyGithub的create_file/update_file接受bytes类型时会自动进行base64编码
        不要手动编码，否则会导致双重编码！
        """
        g, repo_name, branch = GithubSync._get_github_client()
        
        if not g or not repo_name:
            st.error("❌ 未配置 Github Token 或 仓库名")
            return False

        try:
            repo = g.get_repo(repo_name)
            # 注意：直接传bytes，PyGithub会自动base64编码
            # 不要手动编码！否则会导致双重编码，文件损坏
            
            try:
                contents = repo.get_contents(file_path_in_repo, ref=branch)
                repo.update_file(
                    path=contents.path,
                    message=commit_msg,
                    content=file_content,  # 直接传bytes
                    sha=contents.sha,
                    branch=branch
                )
            except GithubException as e:
                if e.status == 404:
                    repo.create_file(
                        path=file_path_in_repo,
                        message=f"Create {file_path_in_repo}",
                        content=file_content,  # 直接传bytes
                        branch=branch
                    )
                else:
                    raise e
            return True

        except Exception as e:
            st.error(f"Github 文件上传失败: {str(e)}")
            return False

    @staticmethod
    def delete_file(file_path_in_repo: str, commit_msg: str = "Delete file") -> bool:
        """从 Github 删除文件"""
        g, repo_name, branch = GithubSync._get_github_client()
        
        if not g or not repo_name:
            return False

        try:
            repo = g.get_repo(repo_name)
            
            try:
                contents = repo.get_contents(file_path_in_repo, ref=branch)
                repo.delete_file(
                    path=contents.path,
                    message=commit_msg,
                    sha=contents.sha,
                    branch=branch
                )
                return True
            except GithubException as e:
                if e.status == 404:
                    return True  # 文件本来就不存在
                raise e

        except Exception as e:
            st.error(f"Github 删除文件失败: {str(e)}")
            return False

    @staticmethod
    def backup_rag_file(file_content: bytes, filename: str, backup_folder: str = "tea_backup") -> bool:
        """
        将RAG文件备份到tea_backup文件夹（只做加法，不删除）
        - file_content: 文件内容bytes
        - filename: 文件名
        - backup_folder: 备份文件夹路径
        返回: 是否成功
        """
        file_path = f"{backup_folder}/{filename}"
        try:
            result = GithubSync.push_binary_file(file_path, file_content, f"Backup RAG file: {filename}")
            if result:
                print(f"[INFO] ✅ 已备份到 {file_path}")
            return result
        except Exception as e:
            print(f"[WARN] 备份文件 {filename} 到 {backup_folder} 失败: {e}")
            return False

    @staticmethod
    def add_rag_files(uploaded_files: List, rag_folder: str = "tea_data/RAG") -> Tuple[bool, List[str]]:
        """
        添加RAG文件到GitHub（只添加，不删除现有文件）
        同时备份到 tea_backup 文件夹（只做加法）
        - uploaded_files: Streamlit上传的文件对象列表
        - rag_folder: GitHub上的RAG文件夹路径
        返回: (是否成功, 成功上传的文件名列表)
        """
        g, repo_name, branch = GithubSync._get_github_client()
        
        if not g or not repo_name:
            st.error("❌ 未配置 Github Token 或 仓库名")
            return False, []

        try:
            uploaded_names = []
            for uf in uploaded_files:
                file_path = f"{rag_folder}/{uf.name}"
                uf.seek(0)
                file_content = uf.read()
                if GithubSync.push_binary_file(file_path, file_content, f"Add RAG file: {uf.name}"):
                    uploaded_names.append(uf.name)
                    # >>> 变更3：同时备份到 tea_backup 文件夹 <<<
                    GithubSync.backup_rag_file(file_content, uf.name, backup_folder="tea_backup")
                else:
                    st.warning(f"⚠️ 上传 {uf.name} 失败")
            
            return len(uploaded_names) > 0, uploaded_names

        except Exception as e:
            st.error(f"RAG文件添加失败: {str(e)}")
            return False, []

    @staticmethod
    def list_rag_files(rag_folder: str = "tea_data/RAG") -> List[str]:
        """
        获取GitHub上RAG文件夹中的所有文件名
        返回: 文件名列表
        """
        g, repo_name, branch = GithubSync._get_github_client()
        
        if not g or not repo_name:
            return []

        try:
            repo = g.get_repo(repo_name)
            contents = repo.get_contents(rag_folder, ref=branch)
            return [c.name for c in contents if c.type == "file"]
        except GithubException as e:
            if e.status == 404:
                return []  # 文件夹不存在
            print(f"[ERROR] 获取RAG文件列表失败: {e}")
            return []
        except Exception as e:
            print(f"[ERROR] 获取RAG文件列表失败: {e}")
            return []

    @staticmethod
    def delete_rag_file(filename: str, rag_folder: str = "tea_data/RAG") -> bool:
        """
        从GitHub删除单个RAG文件（仅从tea_data/RAG删除，tea_backup中保留）
        - filename: 要删除的文件名
        - rag_folder: GitHub上的RAG文件夹路径
        返回: 是否成功
        """
        file_path = f"{rag_folder}/{filename}"
        return GithubSync.delete_file(file_path, f"Delete RAG file: {filename}")

    @staticmethod
    def sync_cases(cases: List[Dict], file_path: str = "tea_data/case.json") -> bool:
        """同步判例库到GitHub"""
        return GithubSync.push_json(file_path, cases, "Update case.json from App")

    @staticmethod
    def pull_rag_folder(rag_folder: str = "tea_data/RAG") -> List[Tuple[str, bytes]]:
        """
        从 GitHub 拉取 RAG 文件夹中的所有文件
        返回: [(文件名, 文件内容bytes), ...]
        
        优化策略：
        1. 优先使用 Raw URL（最可靠，适合大文件）
        2. 备用方案：Git Blob API（仅小于1MB的文件）
        3. 增加完整性验证：对比文件大小
        4. 支持重试机制
        """
        token, repo_name, branch = GithubSync._get_github_config()
        
        if not token or not repo_name:
            print("[WARN] GitHub config not found, skip pulling RAG")
            return []

        def download_with_retry(url, headers, max_retries=3):
            """带重试的下载函数"""
            for attempt in range(1, max_retries + 1):
                try:
                    response = requests.get(url, headers=headers, timeout=180, stream=True)
                    if response.status_code == 200:
                        # 使用 stream=True 分块下载，避免大文件超时
                        content = b''
                        for chunk in response.iter_content(chunk_size=8192):
                            if chunk:
                                content += chunk
                        return content, True
                    else:
                        print(f"[WARN]     尝试 {attempt}/{max_retries}: HTTP {response.status_code}")
                except Exception as e:
                    print(f"[WARN]     尝试 {attempt}/{max_retries}: {e}")
                    if attempt < max_retries:
                        import time
                        time.sleep(2)  # 等待2秒后重试
            return None, False

        try:
            g = Github(auth=Auth.Token(token))
            repo = g.get_repo(repo_name)
            
            files = []
            print(f"[INFO] ========== 开始从 GitHub 拉取 RAG 文件 ==========")
            print(f"[INFO] 仓库: {repo_name}, 分支: {branch}, 文件夹: {rag_folder}")
            
            try:
                contents = repo.get_contents(rag_folder, ref=branch)
                file_list = [c for c in contents if c.type == "file"]
                print(f"[INFO] 发现 {len(file_list)} 个文件")
                
                for idx, content in enumerate(file_list, 1):
                    print(f"\n[INFO] [{idx}/{len(file_list)}] 正在处理: {content.name}")
                    print(f"[INFO]   → 期望大小: {content.size:,} bytes")
                    file_content = None
                    download_method = None
                    
                    # ===== 方法1：Raw URL（优先，最可靠） =====
                    print(f"[INFO]   → 方法1: Raw URL 下载...")
                    raw_url = f"https://raw.githubusercontent.com/{repo_name}/{branch}/{rag_folder}/{content.name}"
                    headers = {"Authorization": f"Bearer {token}"}
                    file_content, success = download_with_retry(raw_url, headers, max_retries=3)
                    
                    if success and file_content:
                        download_method = "Raw URL"
                        print(f"[INFO]   ✓ 下载完成: {len(file_content):,} bytes")
                    
                    # ===== 方法2：Git Blob API（仅用于小文件 <1MB） =====
                    if file_content is None and content.size < 1024 * 1024:  # 1MB
                        try:
                            print(f"[INFO]   → 方法2: Git Blob API...")
                            blob = repo.get_git_blob(content.sha)
                            if blob.encoding == "base64":
                                file_content = base64.b64decode(blob.content)
                                download_method = "Git Blob"
                                print(f"[INFO]   ✓ 下载完成: {len(file_content):,} bytes")
                        except Exception as e:
                            print(f"[WARN]   ✗ Git Blob 失败: {e}")
                    
                    # ===== 方法3：Download URL（兜底） =====
                    if file_content is None and content.download_url:
                        print(f"[INFO]   → 方法3: Download URL...")
                        headers = {
                            "Authorization": f"Bearer {token}",
                            "Accept": "application/vnd.github.v3.raw"
                        }
                        file_content, success = download_with_retry(content.download_url, headers, max_retries=3)
                        if success:
                            download_method = "Download URL"
                            print(f"[INFO]   ✓ 下载完成: {len(file_content):,} bytes")
                    
                    # ===== 验证文件完整性 =====
                    if file_content:
                        actual_size = len(file_content)
                        expected_size = content.size
                        
                        print(f"[INFO]   → 验证完整性...")
                        print(f"[INFO]     期望: {expected_size:,} bytes")
                        print(f"[INFO]     实际: {actual_size:,} bytes")
                        
                        if actual_size == expected_size:
                            files.append((content.name, file_content))
                            print(f"[INFO]   ✅ {content.name} 完整性验证通过 (方法: {download_method})")
                        else:
                            size_diff = abs(actual_size - expected_size)
                            print(f"[ERROR]  ❌ {content.name} 大小不匹配 (差异: {size_diff:,} bytes)")
                            print(f"[ERROR]     文件可能损坏，跳过...")
                    else:
                        print(f"[ERROR]  ❌ {content.name} 所有下载方法均失败")
                            
            except GithubException as e:
                if e.status == 404:
                    print(f"[INFO] RAG 文件夹不存在: {rag_folder}")
                    return []
                print(f"[ERROR] GitHub API 异常: {e}")
                raise e
            
            print(f"\n[INFO] ========== RAG 拉取完成: {len(files)}/{len(file_list)} 个文件验证通过 ==========\n")
            return files

        except Exception as e:
            print(f"[ERROR] 拉取 RAG 文件夹失败: {e}")
            import traceback
            traceback.print_exc()
            return []


# ==========================================
# [SECTION 2] AI 服务 (Embedding & LLM)
# ==========================================

class AliyunEmbedder:
    def __init__(self, api_key):
        self.model_name = "text-embedding-v4"
        dashscope.api_key = api_key 

    def encode(self, texts: List[str]) -> np.ndarray:
        if not texts: return np.zeros((0, 1024), dtype="float32")
        if isinstance(texts, str): texts = [texts]
        try:
            resp = TextEmbedding.call(model=self.model_name, input=texts)
            if resp.status_code == HTTPStatus.OK:
                return np.array([i['embedding'] for i in resp.output['embeddings']]).astype("float32")
        except: pass
        return np.zeros((len(texts), 1024), dtype="float32")

def llm_normalize_user_input(raw_query: str, client: OpenAI) -> str:
    """使用 LLM 对用户输入做语义规范化 / 去噪"""
    system_prompt = (
        """
          A. 角色与目标
          你是"茶评清洗器"。你的任务是从输入文本中提取并输出只与茶评相关的信息，删除无关内容，保持原意与原有表述风格，只能删减不能修改。
          B. 什么算"相关信息"（保留）
          仅保留与以下内容有关的句子/短语：
          茶的基本信息：茶名/品类、产地、年份、工艺、等级、原料、香型等
          干茶/茶汤/叶底：外观、色泽、条索、汤色、叶底描述
          香气与滋味：香气类型、强弱、层次、回甘、生津、涩感、苦感、甜度、醇厚度、喉韵、体感等
          冲泡信息与表现：器具、投茶量、水温、时间、出汤、几泡变化、耐泡度、适饮建议
          主观评价与结论：好喝/一般/缺点/性价比
          C. 什么算"无关信息"（删除）
          删除与茶评无直接关系的内容，例如：
          与茶无关的生活日常、情绪宣泄、社交聊天、段子
          店铺/物流/客服/包装破损/发货慢（除非"包装异味影响茶"这类直接影响品饮）
          广告、价格链接、优惠券、引流话术、品牌吹水（除非是"性价比"且与品饮结论相关）
          与其它产品/话题无关的对比闲聊
          凑字数内容
          D. 输出格式
          只输出清洗后的茶评正文，不要解释、不加标题、不输出"删除了什么"
          如果输入中没有任何茶评相关信息，则输出："无相关茶评信息"
          E. 操作原则
          尽量保留原句；只做删除/少量拼接
          不要补充不存在的细节，不要推测        
          """
    )

    resp = client.chat.completions.create(
        model="deepseek-chat",
        temperature=0,
        timeout=30,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": raw_query}
        ]
    )
    return resp.choices[0].message.content.strip()

# ==========================================
# [SECTION 1.5] External GraphRAG static KB retrieval (via graphrag_retriever.py)
# ==========================================

def _get_graphrag_artifact_dir() -> str:
    """Resolve GraphRAG artifact directory (env overrides local default)."""
    env_dir = os.getenv("GRAPHRAG_ARTIFACT_DIR", "").strip()
    if env_dir:
        return env_dir
    return str(getattr(PATHS, "GRAPHRAG_DIR", PATHS.DATA_DIR / "graphrag_artifacts"))

def _get_graphrag_retriever() -> 'GraphRAGRetriever | None':
    """Lazily load GraphRAGRetriever and cache in session_state."""
    if st.session_state.get("_gr_retriever_loaded", False):
        return st.session_state.get("_gr_retriever_obj", None)

    artifact_dir = _get_graphrag_artifact_dir()
    edges_path = os.path.join(artifact_dir, "graph_edges.jsonl")
    comm_path = os.path.join(artifact_dir, "communities.json")
    nodes_path = os.path.join(artifact_dir, "graph_nodes.json")
    meta_path = os.path.join(artifact_dir, "chunk_meta.jsonl")
    if not (os.path.exists(edges_path) and os.path.exists(comm_path)):
        for fp, fn in [(nodes_path,"graph_nodes.json"),(meta_path,"chunk_meta.jsonl")]:
            if not os.path.exists(fp):
                print(f"[WARN] GraphRAG artifact 缺失: {fn}")
        st.session_state["_gr_retriever_loaded"] = True
        st.session_state["_gr_retriever_obj"] = None
        return None

    try:
        gr = GraphRAGRetriever(artifact_dir=artifact_dir)
        st.session_state["_gr_retriever_loaded"] = True
        st.session_state["_gr_retriever_obj"] = gr
        return gr
    except Exception as e:
        print(f"[WARN] GraphRAGRetriever init failed: {e}")
        st.session_state["_gr_retriever_loaded"] = True
        st.session_state["_gr_retriever_obj"] = None
        return None
def build_graphrag_artifacts(kb_chunks: list, force_rebuild: bool = False,
                            chunk_source_map: dict = None, file_names: list = None) -> bool:
    """
    从已有的 kb_chunks 构建 GraphRAG artifact 文件，并同步到 GitHub。
    [PATCHED v2.1] 新增 chunk_source_map 和 file_names 参数；检查全部4个文件
    """
    artifact_dir = str(PATHS.GRAPHRAG_DIR)
    # >>> PATCHED: 检查全部4个文件，而非仅edges+communities <<<
    required_files = ["graph_edges.jsonl", "communities.json", "graph_nodes.json", "chunk_meta.jsonl"]
    all_exist = all(os.path.exists(os.path.join(artifact_dir, f)) for f in required_files)
    
    if not force_rebuild and all_exist:
        print("[INFO] GraphRAG artifacts 已存在（4个文件齐全），跳过构建")
        return True
    if not force_rebuild and not all_exist:
        missing = [f for f in required_files if not os.path.exists(os.path.join(artifact_dir, f))]
        print(f"[INFO] GraphRAG artifacts 不完整，缺失: {missing}，将重新构建")
    
    if not kb_chunks:
        print("[WARN] kb_chunks 为空，无法构建 GraphRAG artifacts")
        return False
    
    print(f"\n[INFO] ========== 开始构建 GraphRAG Artifacts ==========")
    print(f"[INFO] 共 {len(kb_chunks)} 个文本块待处理")
    
    try:
        # 1. 将 kb_chunks 转为 Chunk 对象（使用 chunk_source_map 提供有意义的来源）
        chunks = []
        for i, text in enumerate(kb_chunks):
            if not text or not text.strip():
                continue
            if chunk_source_map and i in chunk_source_map:
                source = chunk_source_map[i]
            elif file_names:
                source = ",".join(file_names[:3])
            else:
                source = f"kb_chunk_{i}"
            chunks.append(Chunk(
                chunk_id=str(i),
                text=text.strip(),
                source=source,
                tags={"chunk_index": str(i)}
            ))
        
        if not chunks:
            print("[WARN] 有效文本块为0，无法构建 GraphRAG")
            return False
        
        print(f"[INFO] 有效文本块: {len(chunks)}")
        
        # 2. 构建索引（默认使用 RuleBasedExtractor）
        indexer = GraphRAGIndexer()
        indexer.add_chunks(chunks)
        
        node_count = indexer.graph.number_of_nodes()
        edge_count = indexer.graph.number_of_edges()
        print(f"[INFO] 图构建完成: {node_count} 个节点, {edge_count} 条边")
        
        # 3. 发现社区（茶饮领域实体可能较少，降低 min_size 阈值）
        if node_count > 0:
            communities = indexer.build_communities(min_size=2)
            print(f"[INFO] 发现 {len(communities)} 个社区")
        else:
            print("[WARN] 图中无节点，跳过社区发现")
        
        # 4. 保存到本地
        indexer.save(artifact_dir)
        print(f"[INFO] Artifacts 已保存到: {artifact_dir}")
        
        # 5. 验证生成的文件
        for fname in ["graph_edges.jsonl", "communities.json", "graph_nodes.json", "chunk_meta.jsonl"]:
            fpath = os.path.join(artifact_dir, fname)
            if os.path.exists(fpath):
                print(f"[INFO]   ✅ {fname}: {os.path.getsize(fpath):,} bytes")
            else:
                print(f"[WARN]   ❌ {fname}: 未生成")
        
        # 6. 同步到 GitHub
        print("[INFO] 正在同步 artifacts 到 GitHub...")
        github_ok = _push_graphrag_artifacts_to_github(artifact_dir)
        if github_ok:
            print("[INFO] ✅ GraphRAG artifacts 已同步到 GitHub")
        else:
            print("[WARN] ⚠️ GitHub 同步失败，本地文件已生成，不影响本次运行")
        
        # 7. 清除 retriever 缓存，使下次检索重新加载
        st.session_state["_gr_retriever_loaded"] = False
        st.session_state["_gr_retriever_obj"] = None
        
        print(f"[INFO] ========== GraphRAG Artifacts 构建完成 ==========\n")
        return True
        
    except Exception as e:
        print(f"[ERROR] GraphRAG artifacts 构建失败: {e}")
        import traceback
        traceback.print_exc()
        return False


def _push_graphrag_artifacts_to_github(artifact_dir: str) -> bool:
    """将本地 GraphRAG artifact 文件推送到 GitHub
    [PATCHED v2.1] 加入延迟避免 API 速率限制 + 重试机制 + 优先推送 nodes/meta
    """
    github_folder = "tea_data/graphrag_artifacts"
    # >>> PATCHED: 调整推送顺序 — 先推 nodes 和 meta（之前总失败的），再推 edges 和 communities <<<
    files_to_push = ["graph_nodes.json", "chunk_meta.jsonl", "graph_edges.jsonl", "communities.json"]
    
    all_success = True
    for i_file, fname in enumerate(files_to_push):
        local_path = os.path.join(artifact_dir, fname)
        if not os.path.exists(local_path):
            print(f"[WARN]   ⚠️ 本地文件不存在，跳过: {fname}")
            all_success = False
            continue
        
        # 读取本地文件
        try:
            with open(local_path, "rb") as f:
                file_content = f.read()
        except Exception as e:
            print(f"[ERROR]  读取 {fname} 失败: {e}")
            all_success = False
            continue
        
        github_path = f"{github_folder}/{fname}"
        push_ok = False
        
        # >>> PATCHED: 最多重试3次 <<<
        for attempt in range(1, 4):
            try:
                success = GithubSync.push_binary_file(
                    github_path, file_content,
                    f"Update GraphRAG artifact: {fname} (v2.1)"
                )
                if success:
                    print(f"[INFO]   ✅ 已推送: {github_path} ({len(file_content):,} bytes)")
                    push_ok = True
                    break
                else:
                    print(f"[WARN]   推送 {fname} 返回 False (尝试 {attempt}/3)")
            except Exception as e:
                print(f"[WARN]   推送 {fname} 异常 (尝试 {attempt}/3): {e}")
            
            # 重试前等待
            if attempt < 3:
                time.sleep(2)
        
        if not push_ok:
            print(f"[ERROR]  ❌ {fname} 推送最终失败!")
            all_success = False
        
        # 推送之间短暂等待，避免 GitHub API 速率限制
        if i_file < len(files_to_push) - 1:
            time.sleep(0.5)
    
    return all_success

def graphrag_static_kb_context(query_vec: np.ndarray,
                              kb_index: faiss.Index,
                              kb_chunks: List[str],
                              k_num: Optional[int] = None,
                              top_seed: int = 5,
                              hop: int = 1,
                              max_expand: int = 12) -> Tuple[str, List[str]]:
    """Build KB context using (vector seeds) + GraphRAG expansion over a static KB graph."""
    if kb_index is None or getattr(kb_index, "ntotal", 0) <= 0 or not kb_chunks:
        return "（无手册资料）", []

    # Vector seeds
    D, I = kb_index.search(query_vec, max(k_num, top_seed))
    vector_hits: List[Tuple[str, float]] = []
    for dist, idx in zip(D[0].tolist(), I[0].tolist()):
        if idx is None or idx < 0 or idx >= len(kb_chunks):
            continue
        # NOTE: kb_index is IndexFlatL2 -> D is L2 distance (smaller is better).
        # GraphRAGRetriever expects a vector score where larger is better,
        # so we convert distance to a bounded similarity.
        sim = 1.0 / (1.0 + float(dist))
        vector_hits.append((str(idx), sim))

    # Map for retriever
    chunk_text_map = {str(i): kb_chunks[i] for i in range(len(kb_chunks))}

    gr = _get_graphrag_retriever()
    if gr is None:
        # Fallback: classic vector context
        hits = [kb_chunks[int(cid)] for cid, _ in vector_hits[:k_num]]
        ctx = "\n".join([f"- {h[:240].strip()}..." for h in hits]) if hits else "（无手册资料）"
        return ctx, hits

    try:
        # 使用线程超时保护 gr.expand()，避免 GraphRAG 检索 hang 住评分流程
        import concurrent.futures
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(
                gr.expand,
                vector_hits=vector_hits,
                chunk_text_map=chunk_text_map,
                top_seed=top_seed,
                hop=hop,
                max_expand=max_expand,
                w_vec=0.7,
                w_graph=0.3
            )
            try:
                expanded = future.result(timeout=10)  # 最多等10秒
            except concurrent.futures.TimeoutError:
                print(f"[WARN] GraphRAG expand 超时(>10s)，回退到向量检索")
                future.cancel()
                hits = [kb_chunks[int(cid)] for cid, _ in vector_hits[:k_num]]
                ctx = "\n".join([f"- {h[:240].strip()}..." for h in hits]) if hits else "（无手册资料）"
                return ctx, hits
        seeds = expanded.get("seed_chunks", [])
        exp_chunks = expanded.get("expanded_chunks", [])
        comm = expanded.get("community_summaries", [])

        # Extract text from result dicts
        seed_texts = [s.get("text", "") for s in seeds if s.get("text")]
        exp_texts = [c.get("text", "") for c in exp_chunks if c.get("text")]
        comm_texts = [c.get("summary", "") for c in comm if c.get("summary")]

        # Compose: community summaries first (global), then seeds, then expanded
        parts = []
        if comm_texts:
            parts.append("【GraphRAG 社区摘要】\n" + "\n\n".join(comm_texts[:2]))
        if seed_texts:
            parts.append("【向量检索种子片段】\n" + "\n".join([f"- {s[:240].strip()}..." for s in seed_texts[:k_num]]))
        if exp_texts:
            parts.append("【Graph 扩展片段】\n" + "\n".join([f"- {c[:240].strip()}..." for c in exp_texts[:k_num]]))
        ctx = "\n\n".join(parts) if parts else "（无手册资料）"

        hits_texts = []
        for t in (seed_texts + exp_texts):
            if t and t not in hits_texts:
                hits_texts.append(t)
            if len(hits_texts) >= k_num:
                break
        return ctx, hits_texts
    except Exception as e:
        print(f"[WARN] GraphRAG expand failed, fallback to vector-only. err={e}")
        hits = [kb_chunks[int(cid)] for cid, _ in vector_hits[:k_num]]
        ctx = "\n".join([f"- {h[:240].strip()}..." for h in hits]) if hits else "（无手册资料）"
        return ctx, hits

def run_scoring(text: str, kb_res: Tuple, case_res: Tuple, prompt_cfg: Dict, embedder: AliyunEmbedder, client: OpenAI, model_id: str, k_num: int, c_num: int):
    """执行 RAG 检索与 LLM 评分
    
    返回: (scores_dict, kb_hits, found_cases, system_prompt, user_prompt)
    """
    vec = embedder.encode([text]) 
    
    # --- KB (External GraphRAG over static KB) ---
    ctx_txt, hits = graphrag_static_kb_context(
        query_vec=vec,
        kb_index=kb_res[0],
        kb_chunks=kb_res[1],
        k_num=k_num,
        top_seed=max(5, k_num),
        hop=1,
        max_expand=12
    )
    
    case_txt, found_cases = "（无相似判例）", []
    if case_res[0].ntotal > 0:
        _, idx = case_res[0].search(vec, c_num)
        for i in idx[0]:
            if i < len(case_res[1]) and i >= 0:
                c = case_res[1][i]
                found_cases.append(c)
                sc = c.get('scores', {})
                u_sc = sc.get('优雅性',{}).get('score', 0) if isinstance(sc,dict) and '优雅性' in sc else 0
                k_sc = sc.get('苦涩度',{}).get('score', 0) if isinstance(sc,dict) and '苦涩度' in sc else 0
                case_txt += f"\n参考案例: {c['text'][:30]}... -> 优雅性:{u_sc} 苦涩度:{k_sc}"

    text_case = "\n".join([str(x) for x in found_cases])
    sys_p = prompt_cfg.get('system_template', "")
    user_p = prompt_cfg.get('user_template', "")
    user_p = user_p.replace("{product_desc}", text).replace("{context_text}", ctx_txt).replace("{case_text}", text_case)
    
    try:
        resp = client.chat.completions.create(
            model=model_id,
            messages=[{"role":"system", "content":sys_p}, {"role":"user", "content":user_p}],
            response_format={"type": "json_object"},
            temperature=0.3,
            timeout=60
        )
        # >>> 变更1：返回 sys_p 和 user_p 供 Tab1 展示 <<<
        return json.loads(resp.choices[0].message.content), hits, found_cases, sys_p, user_p
    except Exception as e:
        st.error(f"Inference Error: {e}")
        return None, [], [], sys_p, user_p

# ==========================================
# [SECTION 3] 辅助与可视化
# ==========================================

def parse_file_bytes(filename: str, content: bytes) -> str:
    """
    解析文件内容 (从 bytes) - 用于从 GitHub 拉取的文件
    支持格式: .txt, .pdf, .docx
    
    PDF 采用三级策略：缓存 → PyMuPDF文本提取 → 乱码检测 → OCR
    同一个 PDF（按 MD5）永远只 OCR 一次
    """
    try:
        # 1. 处理 TXT 文件
        if filename.lower().endswith('.txt'):
            text = content.decode('utf-8', errors='ignore')
            print(f"[INFO]     → TXT 解析成功: {len(text)} 字符")
            return text

        # 2. 处理 PDF 文件
        elif filename.lower().endswith('.pdf'):
            return _parse_pdf_with_cache(filename, content)

        # 3. 处理 DOCX 文件
        elif filename.lower().endswith('.docx'):
            doc = Document(BytesIO(content))
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            print(f"[INFO]     → DOCX 解析成功: {len(text)} 字符")
            return text

        else:
            print(f"[WARN]     → 不支持的文件格式: {filename}")
            return ""

    except Exception as e:
        print(f"[ERROR]    ✗ 解析 {filename} 失败: {e}")
        import traceback
        traceback.print_exc()
        return ""


def _is_garbled(text: str) -> bool:
    """
    检测 PDF 提取的文本是否是 CID 字体乱码。
    
    三种乱码模式：
    1. (cid:xxx) 标记 → CID字体无ToUnicode映射
    2. 犭部首字符密集（犌犅犜犕犲狋犺...）→ 英文被错误映射为汉字
    3. ASCII替代乱码（!"#$%&'() 替代中文）→ 中文被映射为ASCII码位
       GB/T 23776 就是这种！表面看全是英文字母和标点，实际是中文
    """
    if not text or len(text) < 50:
        return True

    # 只统计非空白字符
    non_ws = [c for c in text if not c.isspace()]
    total = len(non_ws)
    if total < 30:
        return True

    # ── 检测1: (cid:xxx) 标记 ──
    cid_count = text.count('(cid:')
    if cid_count / total > 0.01:
        print(f"[WARN]     → 乱码检测: (cid:) 标记 = {cid_count} ({cid_count/total:.1%})")
        return True

    # ── 检测2: 犭部首字符（扩大范围 U+7280~U+739F 覆盖犌犅犜等） ──
    dog_count = sum(1 for c in non_ws if '\u7280' <= c <= '\u739f')
    if dog_count / total > 0.02:
        print(f"[WARN]     → 乱码检测: 犭部首字符 = {dog_count} ({dog_count/total:.1%})")
        return True

    # ── 检测3: CJK 字符占比过低（最关键！捕获 ASCII 替代乱码） ──
    # 正常中文文档 CJK 占比应 > 20%。如果 < 5%，几乎肯定是乱码。
    cjk_count = sum(1 for c in non_ws if '\u4e00' <= c <= '\u9fff')
    cjk_ratio = cjk_count / total

    # 同时统计全角字符（ＧＢ／Ｔ 等），乱码PDF常见
    fullwidth_count = sum(1 for c in non_ws if '\uff01' <= c <= '\uff5e')
    fullwidth_ratio = fullwidth_count / total

    # 统计普通 ASCII 字母+标点（排除数字，因为正常文档也有数字）
    ascii_letters_punct = sum(1 for c in non_ws 
                              if ('!' <= c <= '/' or ':' <= c <= '@' 
                                  or '[' <= c <= '`' or '{' <= c <= '~'
                                  or 'A' <= c <= 'Z' or 'a' <= c <= 'z'))
    ascii_ratio = ascii_letters_punct / total

    # 判断：CJK极少 + ASCII字母标点过多 = 中文被映射成了ASCII
    if cjk_ratio < 0.05 and ascii_ratio > 0.25:
        print(f"[WARN]     → 乱码检测: CJK={cjk_count}({cjk_ratio:.1%}), "
              f"ASCII字母标点={ascii_letters_punct}({ascii_ratio:.1%}), "
              f"全角={fullwidth_count}({fullwidth_ratio:.1%}) → ASCII替代乱码")
        return True

    # 更宽松的检测：CJK < 10% 且 全角字符多（乱码PDF常把半角转全角）
    if cjk_ratio < 0.10 and fullwidth_ratio > 0.05 and ascii_ratio > 0.15:
        print(f"[WARN]     → 乱码检测: CJK={cjk_ratio:.1%}, 全角={fullwidth_ratio:.1%}, "
              f"ASCII={ascii_ratio:.1%} → 疑似乱码")
        return True

    return False


def _parse_pdf_with_cache(filename: str, content: bytes, force_ocr: bool = False) -> str:
    """PDF 解析主入口：缓存 → PyMuPDF文本提取 → 乱码则OCR
    
    Args:
        force_ocr: 跳过缓存和文本提取，直接OCR（用于修复错误缓存）
    """
    print(f"[INFO]     → 开始解析 PDF: {filename} ({len(content):,} bytes)")

    if not content.startswith(b'%PDF'):
        print(f"[ERROR]    → 不是有效的 PDF 文件")
        return ""

    # ━━━ 第1步：查缓存（force_ocr时跳过） ━━━
    if not force_ocr:
        cached = OCRCache.get(content)
        if cached:
            # 验证缓存内容是否也是乱码（修复之前缓存了乱码文本的情况）
            if not _is_garbled(cached):
                return cached
            else:
                print(f"[WARN]     → 缓存内容为乱码，将重新解析")

    # ━━━ 第2步：尝试 PyMuPDF 文本提取（force_ocr时跳过） ━━━
    if not force_ocr:
        print(f"[INFO]     → 尝试 PyMuPDF 文本提取...")
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()

            if text.strip() and not _is_garbled(text):
                print(f"[INFO]     → ✅ 文本提取成功: {len(text):,} 字符")
                OCRCache.put(content, text)
                return text
            else:
                print(f"[WARN]     → 文本提取结果为乱码，切换 OCR")
        except Exception as e:
            print(f"[WARN]     → PyMuPDF 提取失败: {e}")

    # ━━━ 第3步：OCR ━━━
    print(f"[INFO]     → 启动 OCR（Tesseract）...")
    try:
        text = _ocr_pdf(content)
        if text.strip():
            # OCR 后也验证一下质量
            if _is_garbled(text):
                print(f"[WARN]     → ⚠️ OCR 结果疑似乱码，但仍保存（可能是渲染问题）")
                print(f"[WARN]     → 前200字: {text[:200]}")
            else:
                print(f"[INFO]     → ✅ OCR 完成: {len(text):,} 字符")
            OCRCache.put(content, text)
            return text
        else:
            print(f"[WARN]     → OCR 结果为空")
            return ""
    except Exception as e:
        print(f"[ERROR]    → OCR 失败: {e}")
        import traceback
        traceback.print_exc()
        return ""


def _ocr_pdf(content: bytes, dpi: int = 300) -> str:
    """PyMuPDF 渲染 + Tesseract OCR"""
    doc = fitz.open(stream=content, filetype="pdf")
    total = len(doc)
    print(f"[INFO]     → PDF 共 {total} 页, DPI={dpi}")

    all_text = []
    for i, page in enumerate(doc, 1):
        try:
            zoom = dpi / 72.0
            pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            page_text = pytesseract.image_to_string(img, lang='chi_sim+eng', config='--psm 6')
            if page_text.strip():
                all_text.append(page_text)
            if i % 10 == 0 or i == total:
                print(f"[INFO]     → OCR 进度: {i}/{total}")
        except Exception as e:
            print(f"[WARN]     → 第 {i} 页 OCR 失败: {e}")
    doc.close()
    return "\n".join(all_text)
    
def parse_file(uploaded_file) -> str:
    """解析上传文件（Streamlit UploadedFile 对象）"""
    try:
        if uploaded_file.name.endswith('.txt'):
            return uploaded_file.read().decode("utf-8")
        if uploaded_file.name.endswith('.pdf'):
            uploaded_file.seek(0)
            content = uploaded_file.read()
            return _parse_pdf_with_cache(uploaded_file.name, content)
        if uploaded_file.name.endswith('.docx'):
            return "\n".join([p.text for p in Document(uploaded_file).paragraphs])
    except Exception as e:
        print(f"[ERROR] parse_file 失败: {e}")
        return ""
    return ""


def create_word_report(results: List[Dict]) -> BytesIO:
    """生成Word报告"""
    doc = Document()
    doc.add_heading("茶评批量评分报告", 0)
    for item in results:
        doc.add_heading(f"条目 {item['id']}", 1)
        doc.add_paragraph(f"原文：{item['text']}")
        s = item.get('scores', {}).get('scores', {})
        mc = item.get('scores', {}).get('master_comment', '')
        if mc: doc.add_paragraph(f"总评：{mc}", style="Intense Quote")
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = '因子', '分数', '评语', '建议'
        for k, v in s.items():
            r = table.add_row().cells
            r[0].text = k
            r[1].text = str(v.get('score',''))
            r[2].text = v.get('comment','')
            r[3].text = v.get('suggestion','')
        doc.add_paragraph("_"*20)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def plot_flavor_shape(scores_data: Dict):
    """绘制风味形态图"""
    s = scores_data["scores"]
    top = (s["优雅性"]["score"] + s["辨识度"]["score"]) / 2
    mid = (s["协调性"]["score"] + s["饱和度"]["score"]) / 2
    base = (s["持久性"]["score"] + s["苦涩度"]["score"]) / 2
    
    fig, ax = plt.subplots(figsize=(4, 5))
    fig.patch.set_alpha(0); ax.patch.set_alpha(0)

    y = np.array([1, 2, 3]) 
    x = np.array([base, mid, top])
    y_new = np.linspace(1, 3, 300)
    try:
        spl = make_interp_spline(y, x, k=2)
        x_smooth = spl(y_new)
    except:
        x_smooth = np.interp(y_new, y, x)
    x_smooth = np.maximum(x_smooth, 0.1)

    colors = {'base': '#8B4513', 'mid': '#D2691E', 'top': '#FFD700'}
    for mask, col in [((y_new>=1.0)&(y_new<=1.6), colors['base']), 
                      ((y_new>1.6)&(y_new<=2.4), colors['mid']), 
                      ((y_new>2.4)&(y_new<=3.0), colors['top'])]:
        ax.fill_betweenx(y_new[mask], -x_smooth[mask], x_smooth[mask], color=col, alpha=0.9, edgecolor=None)

    ax.plot(x_smooth, y_new, 'k', linewidth=1, alpha=0.2)
    ax.plot(-x_smooth, y_new, 'k', linewidth=1, alpha=0.2)
    ax.axhline(y=1.6, color='w', linestyle=':', alpha=0.5)
    ax.axhline(y=2.4, color='w', linestyle=':', alpha=0.5)
    
    font = {'ha': 'center', 'va': 'center', 'color': 'white', 'fontweight': 'bold', 'fontsize': 12}
    ax.text(0, 2.7, f"Top\n{top:.1f}", **font)
    ax.text(0, 2.0, f"Mid\n{mid:.1f}", **font)
    ax.text(0, 1.3, f"Base\n{base:.1f}", **font)
    ax.axis('off'); ax.set_xlim(-10, 10); ax.set_ylim(0.8, 3.2)
    return fig

def bootstrap_seed_cases(embedder: AliyunEmbedder):
    """初始化判例库：如果内存/磁盘中为空，则从 case.json 文件读取"""
    case_idx, case_data = st.session_state.cases
    if len(case_data) > 0: return

    # 从外部 JSON 加载 (修改路径)
    seed_cases = ResourceManager.load_external_json(PATHS.SRC_CASES)
    if not seed_cases:
        # 兼容旧路径
        old_path = Path("seed_case.json")
        seed_cases = ResourceManager.load_external_json(old_path)
    
    if not seed_cases:
        st.warning("case.json 未找到或为空，判例库初始化跳过。")
        return

    texts = [c["text"] for c in seed_cases]
    vecs = embedder.encode(texts)

    if case_idx.ntotal == 0: case_idx = faiss.IndexFlatL2(1024)
    if len(vecs) > 0:
        case_idx.add(vecs)
        case_data.extend(seed_cases)
        st.session_state.cases = (case_idx, case_data)
        ResourceManager.save(case_idx, case_data, PATHS.case_index, PATHS.case_data, is_json=True)

def load_rag_from_github(aliyun_key: str) -> Tuple[bool, str]:
    """
    从 GitHub 加载 RAG 文件
    返回: (是否成功, 消息)
    """
    print("\n[INFO] ========== 开始从 GitHub 加载 RAG 数据 ==========")
    
    # >>> 新增：先拉取 OCR 缓存，避免重复 OCR <<<
    OCRCache.pull_all_from_github()
    
    try:
        # 1. 拉取文件
        print("[INFO] 步骤 1/4: 从 GitHub 拉取 RAG 文件...")
        rag_files = GithubSync.pull_rag_folder("tea_data/RAG")
        
        if not rag_files:
            msg = "GitHub 上没有找到 RAG 文件，或所有文件下载失败"
            print(f"[WARN] {msg}")
            return False, msg
        
        print(f"[INFO] 成功拉取并验证 {len(rag_files)} 个文件")
        
        # 2. 解析文件内容
        print("[INFO] 步骤 2/4: 解析文件内容...")
        all_text = ""
        file_names = []
        parse_success = 0
        parse_failed = []
        file_text_lengths = []  # 同时记录长度，避免后续二次解析
        
        for fname, fcontent in rag_files:
            file_names.append(fname)
            print(f"\n[INFO]   → 解析 {fname} ({len(fcontent):,} bytes)...")
            
            parsed_text = parse_file_bytes(fname, fcontent)
            if parsed_text and len(parsed_text.strip()) > 100:
                flen = len(parsed_text) + 1
                all_text += parsed_text + "\n"
                parse_success += 1
                print(f"[INFO]   ✅ 成功提取 {len(parsed_text):,} 字符")
            else:
                flen = 0
                parse_failed.append(fname)
                print(f"[WARN]   ❌ 提取失败或文本过短 ({len(parsed_text) if parsed_text else 0} 字符)")
            file_text_lengths.append((fname, flen))
        
        print(f"\n[INFO] 文件解析完成: {parse_success}/{len(rag_files)} 成功")
        if parse_failed:
            print(f"[WARN] 解析失败的文件: {', '.join(parse_failed)}")
        
        if not all_text.strip():
            msg = f"无法从 RAG 文件中提取有效文本（共尝试 {len(rag_files)} 个文件）"
            print(f"[ERROR] {msg}")
            return False, msg
        
        # 3. 切片
        print(f"\n[INFO] 步骤 3/4: 将文本切片...")
        print(f"[INFO]   → 总文本长度: {len(all_text):,} 字符")
        chunks = [all_text[i:i+600] for i in range(0, len(all_text), 500)]
        print(f"[INFO]   ✓ 切片完成: {len(chunks)} 个片段")
        
        # 构建 chunk→源文件 映射（复用 file_text_lengths，无需二次解析）
        chunk_source_map = {}
        if file_names and parse_success > 0:
            try:
                for ci in range(len(chunks)):
                    cs = ci * 500
                    acc = 0
                    matched = False
                    for fn_i, fl_i in file_text_lengths:
                        if fl_i == 0:
                            continue
                        if cs < acc + fl_i:
                            chunk_source_map[ci] = fn_i
                            matched = True
                            break
                        acc += fl_i
                    if not matched:
                        last_valid = [fn_i for fn_i, fl_i in file_text_lengths if fl_i > 0]
                        if last_valid:
                            chunk_source_map[ci] = last_valid[-1]
                print(f"[INFO]   ✓ chunk 来源映射: {len(chunk_source_map)} 条")
            except Exception as e:
                print(f"[WARN]   chunk 来源映射构建失败: {e}")
                chunk_source_map = {}
        
        if not chunks:
            msg = "切片失败"
            print(f"[ERROR] {msg}")
            return False, msg
        
        # 4. 向量化并构建索引
        print("\n[INFO] 步骤 4/4: 向量化并构建 FAISS 索引...")
        temp_embedder = AliyunEmbedder(aliyun_key)
        kb_idx = faiss.IndexFlatL2(1024)
        
        print(f"[INFO]   → 调用阿里云 Embedding API (批次大小: 25)...")
        
        # 分批向量化，避免 API 限制
        batch_size = 25
        all_vecs = []
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i+batch_size]
            try:
                vecs = temp_embedder.encode(batch)
                all_vecs.append(vecs)
                print(f"[INFO]   → 已处理 {min(i+batch_size, len(chunks))}/{len(chunks)} 片段")
            except Exception as e:
                print(f"[WARN]   → 批次 {i}-{i+batch_size} 向量化失败: {e}")
        
        if not all_vecs:
            msg = "向量化失败"
            print(f"[ERROR] {msg}")
            return False, msg
        
        vecs = np.vstack(all_vecs)
        print(f"[INFO]   ✓ 获得向量: {vecs.shape}")
        
        kb_idx.add(vecs)
        print(f"[INFO]   ✓ FAISS 索引构建完成 (共 {kb_idx.ntotal} 条)")
        
        # 5. 保存到 session_state 和磁盘
        st.session_state.kb = (kb_idx, chunks)
        st.session_state.kb_files = file_names
        
        # Build GraphRAG-style community summaries for static KB chunks (non-case)
        ResourceManager.save(kb_idx, chunks, PATHS.kb_index, PATHS.kb_chunks)
        ResourceManager.save_kb_files(file_names)
        
        # 标记需要构建 GraphRAG（延迟到下一轮渲染，不阻塞当前流程）
        st.session_state['_graphrag_build_pending'] = True
        st.session_state['_graphrag_build_chunks'] = chunks
        st.session_state['_graphrag_build_chunk_source_map'] = chunk_source_map
        st.session_state['_graphrag_build_file_names'] = file_names
        
        success_files = [f for f in file_names if f not in parse_failed]
        msg = f"✅ 成功加载 {len(chunks)} 条知识片段\n📁 来源文件: {', '.join(success_files)}"
        if parse_failed:
            msg += f"\n⚠️  解析失败: {', '.join(parse_failed)}"
        
        print(f"[INFO] {msg}")
        print("[INFO] ========== RAG 加载完成 ==========\n")
        return True, msg
        
    except Exception as e:
        msg = f"加载失败: {str(e)}"
        print(f"[ERROR] ❌ {msg}")
        import traceback
        traceback.print_exc()
        print("[INFO] ========== RAG 加载失败 ==========\n")
        return False, msg


# ==========================================
# [SECTION 3.5] 判例管理弹窗
# ==========================================

@st.dialog("📋 判例库管理", width="large")
def show_cases_dialog(embedder: AliyunEmbedder):
    """展示并管理所有判例的弹窗"""
    cases = st.session_state.cases[1]
    
    if not cases:
        st.info("当前判例库为空")
        return
    
    st.write(f"共 **{len(cases)}** 条判例")
    st.caption("💡 勾选要删除的判例，然后点击底部的确认按钮")
    
    # 用于追踪编辑状态
    if 'editing_case_idx' not in st.session_state:
        st.session_state.editing_case_idx = None
    
    # 使用checkbox收集要删除的判例（不会触发rerun导致弹窗关闭）
    selected_to_delete = []
    
    for idx, case in enumerate(cases):
        with st.container(border=True):
            col1, col2, col3 = st.columns([6, 1, 1])
            
            with col1:
                # 显示判例摘要
                text_preview = case.get('text', '')[:100] + ('...' if len(case.get('text', '')) > 100 else '')
                st.markdown(f"**#{idx+1}** {text_preview}")
                
                # 显示分数摘要
                scores = case.get('scores', {})
                if scores:
                    score_str = " | ".join([f"{k}:{v.get('score', '?')}" for k, v in scores.items()])
                    st.caption(score_str)
            
            with col2:
                if st.button("✏️", key=f"edit_{idx}", help="编辑此判例"):
                    st.session_state.editing_case_idx = idx
                    st.rerun()
            
            with col3:
                # 使用checkbox代替button，避免rerun导致弹窗关闭
                if st.checkbox("删除", key=f"del_check_{idx}", label_visibility="collapsed"):
                    selected_to_delete.append(idx)
    
    # 如果有选中要删除的判例
    if selected_to_delete:
        st.warning(f"⚠️ 已选中 {len(selected_to_delete)} 条判例待删除")
        if st.button("✅ 确认删除并同步", type="primary", use_container_width=True):
            # 执行删除
            new_cases = [c for i, c in enumerate(cases) if i not in selected_to_delete]
            
            # 重建FAISS索引
            new_idx = faiss.IndexFlatL2(1024)
            if new_cases:
                texts = [c["text"] for c in new_cases]
                vecs = embedder.encode(texts)
                new_idx.add(vecs)
            
            st.session_state.cases = (new_idx, new_cases)
            ResourceManager.save(new_idx, new_cases, PATHS.case_index, PATHS.case_data, is_json=True)
            
            # 同步到GitHub
            with st.spinner("同步到GitHub..."):
                GithubSync.sync_cases(new_cases)
            
            st.success("删除完成！")
            time.sleep(1)
            st.rerun()


@st.dialog("✏️ 编辑判例", width="large")
def edit_case_dialog(case_idx: int, embedder: AliyunEmbedder):
    """编辑单个判例的弹窗"""
    cases = st.session_state.cases[1]
    if case_idx >= len(cases):
        st.error("判例不存在")
        return
    
    case = cases[case_idx]
    factors = ["优雅性", "辨识度", "协调性", "饱和度", "持久性", "苦涩度"]
    
    st.subheader(f"编辑判例 #{case_idx + 1}")
    
    # 编辑文本
    new_text = st.text_area("判例描述", case.get("text", ""), height=100)
    new_master = st.text_area("总评", case.get("master_comment", ""), height=60)
    new_tags = st.text_input("标签", case.get("tags", ""))
    
    # 编辑各因子分数
    st.markdown("**因子评分**")
    new_scores = {}
    cols = st.columns(3)
    
    old_scores = case.get("scores", {})
    for i, f in enumerate(factors):
        with cols[i % 3]:
            with st.container(border=True):
                st.markdown(f"**{f}**")
                old_f = old_scores.get(f, {})
                new_scores[f] = {
                    "score": st.number_input(f"分数", 0, 9, int(old_f.get("score", 5)), key=f"edit_s_{f}"),
                    "comment": st.text_input(f"评语", old_f.get("comment", ""), key=f"edit_c_{f}"),
                    "suggestion": st.text_input(f"建议", old_f.get("suggestion", ""), key=f"edit_sg_{f}")
                }
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("💾 保存修改并同步", type="primary"):
            # 更新判例
            cases[case_idx] = {
                "text": new_text,
                "scores": new_scores,
                "tags": new_tags,
                "master_comment": new_master,
                "created_at": case.get("created_at", time.strftime("%Y-%m-%d"))
            }
            
            # 重建FAISS索引（因为文本可能变了）
            new_idx = faiss.IndexFlatL2(1024)
            texts = [c["text"] for c in cases]
            vecs = embedder.encode(texts)
            new_idx.add(vecs)
            
            st.session_state.cases = (new_idx, cases)
            ResourceManager.save(new_idx, cases, PATHS.case_index, PATHS.case_data, is_json=True)
            
            # 同步到GitHub
            with st.spinner("同步到GitHub..."):
                GithubSync.sync_cases(cases)
            
            st.session_state.editing_case_idx = None
            st.success("保存成功！")
            time.sleep(1)
            st.rerun()
    
    with col2:
        if st.button("❌ 取消"):
            st.session_state.editing_case_idx = None
            st.rerun()


# ==========================================
# [SECTION 4] 主程序逻辑
# ==========================================
# A. 初始化 Session
if 'loaded' not in st.session_state:
    print("\n" + "="*70)
    print("[INFO] ========== 茶饮六因子AI评分器 - 系统初始化 ==========")
    print("="*70)
    
    # 1. 加载本地缓存的 RAG 与判例数据
    print("[INFO] 步骤 1/3: 加载本地缓存数据...")
    kb_idx, kb_data = ResourceManager.load(PATHS.kb_index, PATHS.kb_chunks)
    case_idx, case_data = ResourceManager.load(PATHS.case_index, PATHS.case_data, is_json=True)
    st.session_state.kb = (kb_idx, kb_data)
    st.session_state.cases = (case_idx, case_data)
    st.session_state.kb_files = ResourceManager.load_kb_files()
    
    print(f"[INFO]   → 知识库: {len(kb_data)} 个片段")
    print(f"[INFO]   → 判例库: {len(case_data)} 条判例")
    print(f"[INFO]   → RAG 文件: {st.session_state.kb_files}")
    
    # 2. 标记是否需要从 GitHub 加载 RAG（延迟加载）
    print("[INFO] 步骤 2/3: 检查 RAG 状态...")
    if len(kb_data) == 0:
        st.session_state.rag_loading_needed = True
        st.session_state.rag_loading_status = "pending"
        print("[INFO]   ⚠️  本地知识库为空，将在应用启动后从 GitHub 加载")
    else:
        st.session_state.rag_loading_needed = False
        st.session_state.rag_loading_status = "complete"
        print(f"[INFO]   ✅ 使用本地缓存: {len(kb_data)} 个片段")
    
    # 3. 加载 Prompt 配置
    print("[INFO] 步骤 3/3: 加载 Prompt 配置...")
    if PATHS.prompt_config_file.exists():
        try:
            with open(PATHS.prompt_config_file, 'r', encoding='utf-8') as f:
                st.session_state.prompt_config = json.load(f)
                print("[INFO]   ✅ 已加载自定义 Prompt 配置")
        except Exception as e:
            print(f"[WARN]   ⚠️  加载失败: {e}，使用默认配置")
        
    if 'prompt_config' not in st.session_state:
        sys_prompt_content = ResourceManager.load_external_text(PATHS.SRC_SYS_PROMPT, fallback="你是一名茶评专家...")
        st.session_state.prompt_config = {
            "system_template": sys_prompt_content,
            "user_template": DEFAULT_USER_TEMPLATE
        }
        print("[INFO]   ✅ 使用默认 Prompt 配置")
    
    st.session_state.loaded = True
    print("="*70)
    print("[INFO] ========== 系统初始化完成（快速启动模式）==========")
    print("="*70 + "\n")



# B. 侧边栏
with st.sidebar:
    st.header("⚙️ 系统配置")
    st.markdown("**🔐 API 配置**")
    aliyun_key = os.getenv("ALIYUN_API_KEY") or st.secrets.get("ALIYUN_API_KEY", "")
    deepseek_key = os.getenv("DEEPSEEK_API_KEY") or st.secrets.get("DEEPSEEK_API_KEY", "")

    if not aliyun_key or not deepseek_key:
        st.warning("⚠️ 未配置 API Key")
        st.stop()
    else:
        st.success("✅ API 就绪")

    st.markdown("---")
    st.markdown(f"**预处理模型：** `Deepseek-chat`")
    st.markdown(f"**评分模型：** `Qwen3-14B`")
    model_id = "Qwen3-14B"
    try:
        resp = requests.get("http://117.50.138.123:8001/status", timeout=2)
        if resp.status_code == 200 and resp.json().get("lora_available"):
            model_id = "default_lora"
            st.success("🎉 已启用微调模型")
    except:
        pass
    ft_status = ResourceManager.load_ft_status()
    if ft_status and ft_status.get("status") == "succeeded":
        st.info(f"🎉 发现微调模型：`{ft_status.get('fine_tuned_model')}`")

    embedder = AliyunEmbedder(aliyun_key)
    client = OpenAI(api_key="dummy", base_url="http://117.50.138.123:8000/v1")
    client_d = OpenAI(api_key=deepseek_key, base_url="https://api.deepseek.com")
    
    bootstrap_seed_cases(embedder)
    
    st.markdown("---")
    
    # ===== 延迟加载 RAG 逻辑 =====
    kb_files = st.session_state.get('kb_files', [])
    kb_count = len(st.session_state.kb[1])
    case_count = len(st.session_state.cases[1])
    
    # 检查是否需要从 GitHub 加载 RAG
    if st.session_state.get('rag_loading_needed', False):
        loading_status = st.session_state.get('rag_loading_status', 'pending')
        
        if loading_status == 'pending':
            # 显示加载状态
            with st.status("🔄 正在从 GitHub 加载知识库...", expanded=True) as status:
                st.write("📥 下载 RAG 文件...")
                st.session_state.rag_loading_status = 'loading'
                
                try:
                    # 执行加载
                    success, msg = load_rag_from_github(aliyun_key)
                    
                    if success:
                        status.update(label="✅ 知识库加载完成", state="complete", expanded=False)
                        st.session_state.rag_loading_status = 'complete'
                        st.session_state.rag_loading_needed = False
                        time.sleep(1)
                        st.rerun()
                    else:
                        status.update(label="❌ 知识库加载失败", state="error", expanded=True)
                        st.error(msg)
                        st.info("💡 您可以在 Tab3 手动上传 RAG 文件")
                        st.session_state.rag_loading_status = 'failed'
                        
                        # 添加重试按钮
                        if st.button("🔄 重试加载", type="secondary"):
                            st.session_state.rag_loading_status = 'pending'
                            st.rerun()
                except Exception as e:
                    status.update(label="❌ 加载出错", state="error", expanded=True)
                    st.error(f"加载失败: {str(e)}")
                    st.session_state.rag_loading_status = 'failed'
                    
                    if st.button("🔄 重试加载", type="secondary"):
                        st.session_state.rag_loading_status = 'pending'
                        st.rerun()
        
        elif loading_status == 'loading':
            st.info("🔄 正在加载知识库，请稍候...")
        
        elif loading_status == 'failed':
            st.warning("⚠️ 知识库加载失败")
            if st.button("🔄 重试从 GitHub 加载", type="secondary"):
                st.session_state.rag_loading_status = 'pending'
                st.rerun()
    
    # 更新显示的数据
    kb_count = len(st.session_state.kb[1])
    kb_files = st.session_state.get('kb_files', [])
    
    # 延迟构建 GraphRAG artifacts（不阻塞 RAG 加载和评分交互）
    if st.session_state.get('_graphrag_build_pending', False):
        try:
            pending_chunks = st.session_state.pop('_graphrag_build_chunks', [])
            pending_map = st.session_state.pop('_graphrag_build_chunk_source_map', {})
            pending_fnames = st.session_state.pop('_graphrag_build_file_names', [])
            st.session_state['_graphrag_build_pending'] = False
            if pending_chunks:
                print("[INFO] 延迟构建 GraphRAG artifacts...")
                build_graphrag_artifacts(
                    pending_chunks, force_rebuild=True,
                    chunk_source_map=pending_map,
                    file_names=pending_fnames
                )
        except Exception as e:
            print(f"[WARN] 延迟 GraphRAG 构建失败（不影响评分）: {e}")
    
    st.markdown(f"知识库: **{kb_count}** 条 | 判例库: **{case_count}** 条")
    if kb_files:
        pass
    elif kb_count == 0:
        st.caption("⚠️ 知识库为空，请上传文件或从云端加载")
    
    st.caption("快速上传仅支持.zip文件格式。")
    st.caption("少量文件上传请至\"知识库设计\"板块。")
    
    if st.button("📤 导出数据"):
        import zipfile, shutil
        temp_dir = Path("./temp_export"); temp_dir.mkdir(exist_ok=True)
        for p in [PATHS.kb_index, PATHS.kb_chunks, PATHS.case_index, PATHS.case_data, PATHS.prompt_config_file]:
            if p.exists(): shutil.copy2(p, temp_dir / p.name)
        zip_path = Path("./rag_export.zip")
        with zipfile.ZipFile(zip_path, 'w') as z:
            for f in temp_dir.iterdir(): z.write(f, f.name)
        with open(zip_path, 'rb') as f:
            st.download_button("⬇️ 下载ZIP", f, "tea_data.zip", "application/zip")
        shutil.rmtree(temp_dir); zip_path.unlink()

    if st.button("📥 导入数据"):
        u_zip = st.file_uploader("上传ZIP", type=['zip'])
        if u_zip:
            import zipfile, tempfile
            with tempfile.TemporaryDirectory() as td:
                zp = Path(td)/"u.zip"
                with open(zp,'wb') as f: f.write(u_zip.getvalue())
                with zipfile.ZipFile(zp,'r') as z: z.extractall(PATHS.DATA_DIR)
                st.success("导入成功，请刷新"); st.rerun()

# C. 主界面
st.markdown('<div class="main-title">🍵 茶品六因子 AI 评分器 Pro</div>', unsafe_allow_html=True)
st.markdown('<div class="slogan">"一片叶子落入水中，改变了水的味道..."</div>', unsafe_allow_html=True)
st.markdown('<p style="text-align:center; color:#888; font-size:0.95em;">推理服务开放时间：9:00~20:00</p>', unsafe_allow_html=True)

tab1, tab2, tab3, tab4, tab5 = st.tabs(["💡 交互评分", "🚀 批量评分", "📕 知识库设计", "🛠️ 判例库与微调", "📲 提示词（Prompt）配置"])

# --- Tab 1: 交互评分 ---
with tab1:
    st.info("将参考知识库与判例库进行评分。确认结果可一键更新判例库。")
    c1, c2, c3, c4, c5 = st.columns([1, 3, 1, 3, 1])
    r_num = c2.number_input("参考知识库条目数量", 1, 20, 3, key="r1")
    c_num = c4.number_input("参考判例库条目数量", 1, 20, 2, key="c1")
    
    if 'current_user_input' not in st.session_state: st.session_state.current_user_input = ""
    user_input = st.text_area("请输入茶评描述:", value=st.session_state.current_user_input, height=150, key="ui")
    st.session_state.current_user_input = user_input
    
    if 'last_scores' not in st.session_state: 
        st.session_state.last_scores = None
        st.session_state.last_master_comment = ""
    
    # >>> 变更1：初始化用于存储发送给LLM的prompt的session_state <<<
    if 'last_llm_sys_prompt' not in st.session_state:
        st.session_state.last_llm_sys_prompt = ""
    if 'last_llm_user_prompt' not in st.session_state:
        st.session_state.last_llm_user_prompt = ""
    
    # 用于生成动态key，确保每次新评分时校准输入框显示新内容
    if 'score_version' not in st.session_state:
        st.session_state.score_version = 0
    
    if st.button("开始评分", type="primary", use_container_width=True):
        if not user_input: st.warning("请输入内容")
        else:
            with st.spinner(f"正在使用 {model_id} 品鉴..."):
                try:
                    # 步骤1：预处理（DeepSeek去噪）
                    try:
                        user_input = llm_normalize_user_input(user_input, client_d)
                    except Exception as e:
                        st.warning(f"⚠️ 预处理超时或失败({e})，将使用原始输入继续评分")
                    
                    # 步骤2：RAG检索 + LLM评分
                    scores, kb_h, case_h, sent_sys_p, sent_user_p = run_scoring(
                        user_input, st.session_state.kb, st.session_state.cases,
                        st.session_state.prompt_config, embedder, client, "Qwen3-14B", r_num, c_num
                    )
                    if scores:
                        st.session_state.last_scores = scores
                        st.session_state.last_master_comment = scores.get("master_comment", "")
                        st.session_state.last_llm_sys_prompt = sent_sys_p
                        st.session_state.last_llm_user_prompt = sent_user_p
                        
                        # 递增版本号，使校准输入框使用新的key，从而显示新的默认值
                        st.session_state.score_version += 1
                        st.rerun()
                    else:
                        st.error("❌ 评分失败：模型未返回有效结果。请检查推理服务是否在线（9:00~20:00），或稍后重试。")
                except Exception as e:
                    st.error(f"❌ 评分过程出错：{e}")
    
    if st.session_state.last_scores:
        s = st.session_state.last_scores["scores"]
        mc = st.session_state.last_master_comment
        st.markdown(f'<div class="master-comment"><b>👵 宗师总评：</b><br>{mc}</div>', unsafe_allow_html=True)
        
        left_col, right_col = st.columns([35, 65]) 
        with left_col:
            st.subheader("📊 风味形态")
            st.pyplot(plot_flavor_shape(st.session_state.last_scores), use_container_width=True)
        with right_col:
            cols = st.columns(2)
            factors = ["优雅性", "辨识度", "协调性", "饱和度", "持久性", "苦涩度"]
            for i, f in enumerate(factors):
                if f in s:
                    d = s[f]
                    with cols[i%2]:
                        st.markdown(f"""<div class="factor-card"><div class="score-header"><span>{f}</span><span>{d['score']}/9</span></div><div>{d['comment']}</div><div class="advice-tag">💡 {d.get('suggestion','')}</div></div>""", unsafe_allow_html=True)
        
        # >>> 变更1：展示发送给LLM的Prompt <<<
        with st.expander("📝 查看发送给LLM的完整Prompt", expanded=False):
            st.markdown("**🔧 System Prompt（系统提示词）：**")
            st.code(st.session_state.last_llm_sys_prompt, language=None)
            st.markdown("**💬 User Prompt（用户提示词）：**")
            st.code(st.session_state.last_llm_user_prompt, language=None)
        
        st.subheader("🛠️ 评分校准与修正")
        v = st.session_state.score_version  # 获取当前版本号
        cal_master = st.text_area("校准总评", mc, key=f"cal_master_{v}")
        cal_scores = {}
        st.write("分项调整")
        active_factors = [f for f in factors if f in s]
        grid_cols = st.columns(3) 
        for i, f in enumerate(active_factors):
            with grid_cols[i % 3]:
                with st.container(border=True):
                    t_col, s_col = st.columns([1, 1])
                    with t_col:
                        st.markdown(f"<div style='padding-top: 5px;'><b>📌 {f}</b></div>", unsafe_allow_html=True)
                    with s_col:
                        new_score = st.number_input("分数", 0, 9, int(s[f]['score']), 1, key=f"s_{f}_{v}", label_visibility="collapsed")
                    cal_scores[f] = {
                        "score": new_score,
                        "comment": st.text_area(f"评语", s[f]['comment'], key=f"c_{f}_{v}", height=80, placeholder="评语"),
                        "suggestion": st.text_area(f"建议", s[f].get('suggestion',''), key=f"sg_{f}_{v}", height=68, placeholder="建议")
                    }
        
        if st.button("💾 保存校准评分", type="primary"):
            nc = {"text": user_input, "scores": cal_scores, "tags": "交互-校准", "master_comment": cal_master, "created_at": time.strftime("%Y-%m-%d")}
            st.session_state.cases[1].append(nc)
            st.session_state.cases[0].add(embedder.encode([user_input]))
            ResourceManager.save(st.session_state.cases[0], st.session_state.cases[1], PATHS.case_index, PATHS.case_data, is_json=True)
       
            with st.spinner("同步判例到GitHub..."):
                GithubSync.sync_cases(st.session_state.cases[1])
            
            st.success("校准已保存并同步"); st.rerun()

# --- Tab 2: 批量评分 ---
with tab2:
    f = st.file_uploader("上传文件 (.txt/.docx)")
    c1, c2, c3, c4, c5 = st.columns([1, 3, 1, 3, 1])
    r_n = c2.number_input("参考知识库条目数量", 1, 20, 3, key="rb")
    c_n = c4.number_input("参考判例库条目数量", 1, 20, 2, key="cb")
    if f and st.button("批量处理"):
        lines = [l.strip() for l in parse_file(f).split('\n') if len(l)>10]
        res, bar = [], st.progress(0)
        for i, l in enumerate(lines):
            l = llm_normalize_user_input(l, client_d)
            # run_scoring 现在返回5个值，批量模式忽略后两个prompt
            s, _, _, _, _ = run_scoring(l, st.session_state.kb, st.session_state.cases, st.session_state.prompt_config, embedder, client, "Qwen3-14B", r_n, c_n)
            res.append({"id":i+1, "text":l, "scores":s})
            bar.progress((i+1)/len(lines))
        st.success("完成")
        st.download_button("下载Word", create_word_report(res), "report.docx")

# --- Tab 3: RAG ---
# >>> 变更2：移除Column 2，将"从云端加载知识库"合并到操作流程中（添加/删除后自动重建） <<<
with tab3:
    st.subheader("📚 知识库 (RAG)")
    st.caption("上传PDF/文档以增强模型回答的准确性。文件将同步到云端。添加或删除文件后，系统将自动从云端重建本地知识库。")

    # ===== 显示GitHub上的RAG文件列表 =====
    st.markdown("**📁 云端上的RAG文件：**")
    
    # 获取GitHub上的文件列表
    if 'github_rag_files' not in st.session_state:
        st.session_state.github_rag_files = []
    
    col_refresh, col_spacer = st.columns([1, 3])
    with col_refresh:
        if st.button("🔄 刷新列表", key="refresh_rag_list"):
            with st.spinner("正在获取文件列表..."):
                st.session_state.github_rag_files = GithubSync.list_rag_files()
            st.rerun()
    
    github_files = st.session_state.github_rag_files
    if not github_files:
        # 首次加载时尝试获取
        github_files = GithubSync.list_rag_files()
        st.session_state.github_rag_files = github_files
    
    if github_files:
        st.info(f"共 {len(github_files)} 个文件")
        
        # 用于追踪需要删除的文件
        if 'rag_files_to_delete' not in st.session_state:
            st.session_state.rag_files_to_delete = set()
        
        # 显示文件列表，每个文件带删除按钮
        for fname in github_files:
            file_col, del_col = st.columns([5, 1])
            with file_col:
                if fname in st.session_state.rag_files_to_delete:
                    st.markdown(f"~~📄 {fname}~~ *(待删除)*")
                else:
                    st.markdown(f"📄 {fname}")
            with del_col:
                if fname not in st.session_state.rag_files_to_delete:
                    if st.button("🗑️", key=f"del_rag_{fname}", help=f"删除 {fname}"):
                        st.session_state.rag_files_to_delete.add(fname)
                        st.rerun()
                else:
                    if st.button("↩️", key=f"undo_rag_{fname}", help="撤销删除"):
                        st.session_state.rag_files_to_delete.discard(fname)
                        st.rerun()
        
        # 如果有待删除的文件，显示确认按钮
        if st.session_state.rag_files_to_delete:
            st.warning(f"⚠️ 将删除 {len(st.session_state.rag_files_to_delete)} 个文件")
            del_col1, del_col2 = st.columns(2)
            with del_col1:
                if st.button("✅ 确认删除并同步知识库", type="primary", key="confirm_del_rag"):
                    with st.spinner("正在删除文件..."):
                        deleted = []
                        for fname in st.session_state.rag_files_to_delete:
                            if GithubSync.delete_rag_file(fname):
                                deleted.append(fname)
                        
                        # 更新session state
                        st.session_state.github_rag_files = [f for f in github_files if f not in deleted]
                        
                        # 更新本地知识库文件列表
                        current_kb_files = st.session_state.get('kb_files', [])
                        st.session_state.kb_files = [f for f in current_kb_files if f not in deleted]
                        ResourceManager.save_kb_files(st.session_state.kb_files)
                        
                        st.session_state.rag_files_to_delete = set()
                        st.success(f"✅ 已删除 {len(deleted)} 个文件")
                    
                    # >>> 变更2：删除后自动从云端重建知识库 <<<
                    with st.spinner("🔄 正在从云端重建本地知识库..."):
                        success, msg = load_rag_from_github(aliyun_key)
                        if success:
                            st.success(msg)
                            st.session_state.github_rag_files = GithubSync.list_rag_files()
                        else:
                            st.warning(f"知识库重建失败: {msg}")
                    time.sleep(1)
                    st.rerun()
            with del_col2:
                if st.button("❌ 取消", key="cancel_del_rag"):
                    st.session_state.rag_files_to_delete = set()
                    st.rerun()
    else:
        st.caption("暂无RAG文件")
    
    st.markdown("---")
    
    # ===== 上传新文件（添加模式） =====
    st.markdown("**➕ 添加新文件：**")
    up = st.file_uploader("选择文件", accept_multiple_files=True, key="kb_uploader", 
                        type=['pdf', 'txt', 'docx'])
    
    if up and st.button("📤 添加到知识库并同步", type="primary"):
        # 检查是否有重名文件
        new_names = [u.name for u in up]
        existing_names = st.session_state.get('github_rag_files', [])
        duplicate_names = set(new_names) & set(existing_names)
        
        if duplicate_names:
            st.warning(f"⚠️ 以下文件已存在，将被覆盖：{', '.join(duplicate_names)}")
        
        with st.spinner("正在处理文件..."):
            # 1. 解析文件内容（仅做验证，确保文件可解析）
            raw = "".join([parse_file(u) for u in up])
            
            if not raw.strip():
                st.error("❌ 无法从上传的文件中提取有效文本")
            else:
                # 2. 上传到GitHub（包含tea_data/RAG和tea_backup的备份）
                with st.spinner("上传到GitHub..."):
                    success, uploaded_names = GithubSync.add_rag_files(up, "tea_data/RAG")
                
                if success:
                    # 3. 更新本地文件列表
                    current_kb_files = st.session_state.get('kb_files', [])
                    # 合并文件列表（去重）
                    all_files = list(set(current_kb_files + uploaded_names))
                    st.session_state.kb_files = all_files
                    st.session_state.github_rag_files = list(set(existing_names + uploaded_names))
                    ResourceManager.save_kb_files(all_files)
                    
                    st.success(f"✅ 已上传 {len(uploaded_names)} 个文件到GitHub")
                    
                    # >>> 变更2：上传后自动从云端重建知识库 <<<
                    with st.spinner("🔄 正在从云端重建本地知识库..."):
                        rebuild_success, rebuild_msg = load_rag_from_github(aliyun_key)
                        if rebuild_success:
                            st.success(rebuild_msg)
                            st.session_state.github_rag_files = GithubSync.list_rag_files()
                        else:
                            st.warning(f"知识库重建失败: {rebuild_msg}")
                    time.sleep(1)
                    st.rerun()
                else:
                    st.error("❌ 上传到GitHub失败")
    
    st.markdown("---")
    
    # ===== 手动重建按钮（兜底） =====
    st.markdown("**🔧 手动维护：**")
    st.caption("如果自动重建未生效，可以手动触发从云端重建知识库。")
    if st.button("🔄 手动从云端重建知识库", key="manual_rebuild_kb"):
        with st.spinner("正在从云端拉取并重建知识库..."):
            success, msg = load_rag_from_github(aliyun_key)
            if success:
                st.success(msg)
                st.session_state.github_rag_files = GithubSync.list_rag_files()
            else:
                st.error(msg)
        time.sleep(1)
        st.rerun()


with tab4:
    MANAGER_URL = "http://117.50.138.123:8001"
    c1, c2 = st.columns([5, 5])
    
    with c1:
        st.subheader("📕 判例库 (CASE)")        
        if st.button("📋 展示当前判例", use_container_width=True):
            show_cases_dialog(embedder)
        
        # 检查是否需要打开编辑弹窗
        if st.session_state.get('editing_case_idx') is not None:
            edit_case_dialog(st.session_state.editing_case_idx, embedder)
        
        with st.expander("➕ 手动添加精细判例"):
            with st.form("case_form"):
                f_txt = st.text_area("判例描述", height=80)
                f_tag = st.text_input("标签", "人工录入")
                st.markdown("**因子评分详情**")
                fc1, fc2 = st.columns(2)
                factors = ["优雅性", "辨识度", "协调性", "饱和度", "持久性", "苦涩度"]
                input_scores = {}
                for i, f in enumerate(factors):
                    with (fc1 if i%2==0 else fc2):
                        val = st.number_input(f"{f}分数", 0,9,7, key=f"s_{i}")
                        cmt = st.text_input(f"{f}评语", key=f"c_{i}")
                        sug = st.text_input(f"{f}建议", key=f"a_{i}")
                        input_scores[f] = {"score": val, "comment": cmt, "suggestion": sug}
                
                if st.form_submit_button("保存判例并同步"):
                    new_c = {"text": f_txt, "tags": f_tag, "scores": input_scores, "created_at": time.strftime("%Y-%m-%d")}
                    st.session_state.cases[1].append(new_c)
                    vec = embedder.encode([f_txt])
                    st.session_state.cases[0].add(vec)
                    ResourceManager.save(st.session_state.cases[0], st.session_state.cases[1], PATHS.case_index, PATHS.case_data, is_json=True)
                    
                    # 同步到GitHub
                    with st.spinner("同步到GitHub..."):
                        GithubSync.sync_cases(st.session_state.cases[1])
                    
                    st.success("已保存并同步！")
                    time.sleep(1); st.rerun()

    # --- 右侧：微调控制 ---
    with c2:
        st.subheader("🚀 模型微调 (LoRA)")
        
        admin_pwd = st.text_input("🔒 请输入管理员密码以解锁微调功能", type="password", key="admin_pwd_tab4")
        
        if admin_pwd != "tea_agent_2026":
            if admin_pwd:
                st.error("❌ 密码错误，请重试。")
            else:
                st.info("🔐 输入管理员密码后解锁微调功能。")
        else:
            st.success("✅ 已解锁微调功能")
            
            server_status = "unknown"
            try:
                resp = requests.get(f"{MANAGER_URL}/status", timeout=2)
                if resp.status_code == 200:
                    status_data = resp.json()
                    if status_data.get("vllm_status") == "running":
                        server_status = "idle"
                    else:
                        server_status = "training"
                else:
                    server_status = "error"
            except:
                server_status = "offline"
            
            if server_status == "idle":
                st.success("🟢 服务器就绪 (正在进行推理服务)")
            elif server_status == "training":
                st.warning("🟠 正在微调训练中... (推理服务暂停)")
                st.markdown("⚠️ **注意：** 此时无法进行评分交互，请耐心等待训练完成。")
            elif server_status == "offline":
                st.error("🔴 无法连接到 GPU 服务器 (请联系管理员)")

            st.markdown("#### 1. 数据准备")
            
            if PATHS.training_file.exists():
                with open(PATHS.training_file, "r", encoding="utf-8") as f:
                    lines = f.readlines()
                data_count = len(lines)
            else:
                data_count = 0
                
            st.info(f"当前微调数据：**{data_count} 条** | 判例库：**{len(st.session_state.cases[1])} 条**")
            
            # ===== 修改：覆盖逻辑 =====
            if st.button("🔄 将当前所有判例转为微调数据（覆盖）"):
                cnt = ResourceManager.overwrite_finetune(
                    st.session_state.cases[1],
                    st.session_state.prompt_config.get('system_template',''), 
                    st.session_state.prompt_config.get('user_template','')
                )
                st.success(f"已覆盖写入 {cnt} 条微调数据！")
                time.sleep(1); st.rerun()

            st.markdown("#### 2. 启动训练")
            st.caption("点击下方按钮将把数据上传至 GPU 服务器并开始训练。训练期间服务将中断约 2-5 分钟。")

            btn_disabled = (server_status != "idle") or (data_count == 0)
            
            if st.button("🔥 开始微调 (Start LoRA)", type="primary", disabled=btn_disabled):
                if not PATHS.training_file.exists():
                    st.error("找不到训练数据文件！")
                else:
                    try:
                        with open(PATHS.training_file, "rb") as f:
                            with st.spinner("正在上传数据并启动训练任务..."):
                                files = {'file': ('tea_feedback.jsonl', f, 'application/json')}
                                r = requests.post(f"{MANAGER_URL}/upload_and_train", files=files, timeout=100)
                                
                            if r.status_code == 200:
                                st.balloons()
                                st.success(f"✅ 任务已提交！服务器响应: {r.json().get('message')}")
                                st.info("💡 你可以稍后刷新页面查看状态，训练完成后服务会自动恢复。")
                            else:
                                st.error(f"❌ 提交失败: {r.text}")
                    except Exception as e:
                        st.error(f"❌ 连接错误: {e}")

# --- Tab 5: Prompt配置 ---
with tab5:
    pc = st.session_state.prompt_config
    st.markdown("系统提示词**可以修改**。完整全面的提示词会让大语言模型返回的更准确结果。")    
    sys_t = st.text_area("系统提示词", pc.get('system_template',''), height=350)
    st.markdown("用户提示词**不可修改**。其保证了发送内容与回答内容的基本结构，因此大语言模型的回答可被准确解析。")
    user_t = st.text_area("用户提示词", pc.get('user_template',''), height=250, disabled=True)
    
    if st.button("💾 保存（永久化同步）", type="primary"):
        if sys_t == pc.get('system_template'):
            st.info("内容没有变化，无需保存。")
        else:
            new_cfg = {"system_template": sys_t, "user_template": user_t}
            
            with st.spinner("正在连接云端仓库并写入数据..."):
                success = GithubSync.push_json(
                    file_path_in_repo="tea_data/prompts.json", 
                    data_dict=new_cfg,
                    commit_msg="Update prompts.json from App"
                )
            
            if success:
                st.success("✅ 成功写入云端！")
                st.session_state.prompt_config = new_cfg
                with open(PATHS.prompt_config_file, 'w', encoding='utf-8') as f:
                    json.dump(new_cfg, f, ensure_ascii=False, indent=2)

